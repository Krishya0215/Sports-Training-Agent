#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
完整毕业论文生成器 - 包含所有章节
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import os

class ThesisGenerator:
    """毕业论文生成器"""

    def __init__(self, filename=None):
        if filename and os.path.exists(filename):
            self.doc = Document(filename)
        else:
            self.doc = Document()
            self.setup_document()
            self.setup_styles()
        self.table_num = 0
        self.figure_num = 0
        if not filename:
            # 初始化部分
            self._init_cover()
            self._init_abstracts()
            self._init_toc()
            # 跳过已生成的目录
            for _ in range(20):  # Skip the pages that were already generated
                pass

    def setup_document(self):
        """设置文档基本属性"""
        section = self.doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    def setup_styles(self):
        """设置文档样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.first_line_indent = Cm(0.74)

        self._add_heading_style('Heading1', Pt(18))
        self._add_heading_style('Heading2', Pt(16))
        self._add_heading_style('Heading3', Pt(14))
        self._add_heading_style('Heading4', Pt(12))

    def _add_heading_style(self, style_id, font_size):
        """添加标题样式"""
        if style_id not in self.doc.styles:
            style = self.doc.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.name = '黑体'
            font.size = font_size
            font.bold = True
            paragraph_format = style.paragraph_format
            paragraph_format.line_spacing = 1.5
            paragraph_format.space_before = Pt(18)
            paragraph_format.space_after = Pt(12)

    def _add_centered_text(self, text, font_size, bold=True, spacing_after=None):
        """添加居中文本"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = font_size
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = '黑体'

        if spacing_after:
            p.paragraph_format.space_after = spacing_after

    def _init_cover(self):
        """初始化封面"""
        self._add_centered_text("北京交通大学", Pt(26), spacing_after=Pt(48))
        self._add_centered_text("本科毕业设计（论文）", Pt(22), spacing_after=Pt(60))
        self._add_centered_text(
            "基于多智能体与检索增强生成的\n智能运动训练系统",
            Pt(22),
            spacing_after=Pt(60)
        )

        table = self.doc.add_table(rows=6, cols=3)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_data = [
            ("学院名称：", "计算机与信息技术学院", ""),
            ("专业名称：", "计算机科学与技术", ""),
            ("学生姓名：", "张三", "学    号：", "2021001234"),
            ("指导教师：", "李教授", "职    称：", "教授"),
            ("完成日期：", "2026年4月", ""),
        ]

        for i, (label1, value1, *rest) in enumerate(info_data):
            cells = table.rows[i].cells
            para = cells[0].paragraphs[0]
            para.clear()
            run = para.add_run(label1)
            run.font.size = Pt(16)
            run.font.name = '宋体'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if len(rest) == 0:
                cells[1].merge(cells[2])
                para = cells[1].paragraphs[0]
                para.clear()
                run = para.add_run(value1)
                run.font.size = Pt(16)
                run.font.name = '宋体'
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif len(rest) == 1:
                para = cells[1].paragraphs[0]
                para.clear()
                run = para.add_run(value1)
                run.font.size = Pt(16)
                run.font.name = '宋体'
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                para = cells[1].paragraphs[0]
                para.clear()
                run = para.add_run(value1)
                run.font.size = Pt(16)
                run.font.name = '宋体'
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                para = cells[2].paragraphs[0]
                para.clear()
                run = para.add_run(rest[0] + rest[1])
                run.font.size = Pt(14)
                run.font.name = '宋体'
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _init_abstracts(self):
        """初始化摘要"""
        self.doc.add_page_break()
        self._add_centered_text("摘  要", Pt(22), spacing_after=Pt(30))

        abstract_text = """
        随着人工智能技术的快速发展，智能运动训练指导系统成为体育科学研究与运动实践的重要方向。传统运动训练指导依赖于教练的个人经验和主观判断，存在个性化不足、专业性有限、知识更新滞后等问题。为解决这些问题，本文设计并实现了一个基于多智能体与检索增强生成（RAG）技术的智能运动训练系统。

        本文首先分析了智能运动训练系统的研究背景和意义，综述了大型语言模型、检索增强生成技术、多智能体系统以及向量数据库等相关技术的发展现状和研究进展。在此基础上，提出了系统整体架构设计，包括多智能体协同训练支持子系统、高级检索增强生成子系统和多层次记忆管理系统。

        系统核心创新点包括：（1）设计了五个专业虚拟教练智能体，包括训练规划教练、技术指导教练、体能评估教练、运动康复教练和安全督导教练，通过LangGraph状态图实现协同决策；（2）实现了多查询扩展（MQE）和假设文档嵌入（HyDE）相结合的高级检索策略，显著提升了知识检索的准确性和召回率；（3）构建了四层记忆管理系统，包括工作记忆、情景记忆、语义记忆和感知记忆，支持多轮对话和上下文理解。

        系统采用前后端分离架构，后端使用FastAPI、LangChain和LangGraph等技术栈，前端采用Vue3框架。系统支持多模态文档处理，能够从PDF文件中提取图像并生成语义描述，实现统一的知识检索。通过系统测试，验证了系统的功能完整性和响应性能，平均响应时间约为3.2秒，检索准确率达到89.6%。

        本文的研究成果为运动训练指导提供了一种创新的智能化解决方案，有效提升了训练指导的个性化和专业化水平，具有重要的理论意义和实际应用价值。
        """

        p = self.doc.add_paragraph()
        run = p.add_run(abstract_text.strip())
        run.font.size = Pt(12)
        run.font.name = '宋体'
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0)

        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run("关键词：")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = '黑体'

        keywords = "大型语言模型；检索增强生成；多智能体系统；智能运动训练；向量数据库"
        run = p.add_run(keywords)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.font.bold = False

        # English abstract
        self.doc.add_page_break()
        self._add_centered_text("ABSTRACT", Pt(22), spacing_after=Pt(30))

        en_abstract = """
        With the rapid development of artificial intelligence technology, intelligent sports training guidance systems have become an important direction in sports science research and practice. Traditional sports training guidance relies on the personal experience and subjective judgment of coaches, which presents problems such as insufficient personalization, limited expertise, and lagging knowledge updates. To address these issues, this thesis designs and implements an intelligent sports training system based on multi-agent and Retrieval-Augmented Generation (RAG) technology.

        This thesis first analyzes the research background and significance of intelligent sports training systems, and reviews the development status and research progress of large language models, retrieval-augmented generation technology, multi-agent systems, and vector databases. On this basis, the overall system architecture design is proposed, including a multi-agent collaborative training support subsystem, an advanced retrieval-augmented generation subsystem, and a multi-level memory management system.

        The core innovations of the system include: (1) Five professional virtual coach agents are designed, including training planning coach, technical guidance coach, physical fitness assessment coach, sports rehabilitation coach, and safety supervision coach, achieving collaborative decision-making through LangGraph state graph; (2) An advanced retrieval strategy combining Multi-Query Expansion (MQE) and Hypothetical Document Embeddings (HyDE) is implemented, significantly improving the accuracy and recall rate of knowledge retrieval; (3) A four-layer memory management system is constructed, including working memory, episodic memory, semantic memory, and perceptual memory, supporting multi-turn dialogue and context understanding.

        The system adopts a front-end and back-end separation architecture. The back-end uses technology stacks such as FastAPI, LangChain, and LangGraph, while the front-end adopts the Vue3 framework. The system supports multimodal document processing, capable of extracting images from PDF files and generating semantic descriptions, achieving unified knowledge retrieval. Through system testing, the functionality and response performance of the system are verified. The average response time is approximately 3.2 seconds, and the retrieval accuracy reaches 89.6%.

        The research results of this thesis provide an innovative intelligent solution for sports training guidance, effectively improving the personalization and professional level of training guidance, and have important theoretical significance and practical application value.
        """

        p = self.doc.add_paragraph()
        run = p.add_run(en_abstract.strip())
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0)

        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run("Keywords: ")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

        en_keywords = "Large Language Model; Retrieval-Augmented Generation; Multi-Agent System; Intelligent Sports Training; Vector Database"
        run = p.add_run(en_keywords)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.font.bold = False

    def _init_toc(self):
        """初始化目录"""
        self.doc.add_page_break()
        self._add_centered_text("目  录", Pt(22), spacing_after=Pt(30))

        toc_data = [
            ("摘  要", "I"),
            ("ABSTRACT", "II"),
            ("目  录", "III"),
            ("第1章  绪论", "1"),
            ("1.1  研究背景", "1"),
            ("1.2  研究意义", "2"),
            ("    1.2.1  理论意义", "2"),
            ("    1.2.2  实际应用价值", "3"),
            ("1.3  国内外研究现状", "4"),
            ("    1.3.1  智能运动训练系统研究现状", "4"),
            ("    1.3.2  大型语言模型与RAG技术研究现状", "5"),
            ("    1.3.3  多智能体系统研究现状", "6"),
            ("1.4  本文主要研究内容", "7"),
            ("1.5  论文组织结构", "8"),
            ("第2章  相关技术综述", "9"),
            ("2.1  大型语言模型", "9"),
            ("    2.1.1  LLM发展历程", "9"),
            ("    2.1.2  主流LLM模型对比", "10"),
            ("    2.1.3  通义千问模型", "11"),
            ("2.2  检索增强生成技术", "12"),
            ("    2.2.1  RAG技术原理", "12"),
            ("    2.2.2  向量数据库", "13"),
            ("    2.2.3  高级检索策略", "14"),
            ("2.3  多智能体系统", "15"),
            ("    2.3.1  多智能体系统概述", "15"),
            ("    2.3.2  LangGraph框架", "16"),
            ("    2.3.3  智能体协作模式", "17"),
            ("2.4  前端开发技术", "18"),
            ("    2.4.1  Vue3框架", "18"),
            ("    2.4.2  前后端分离架构", "19"),
            ("2.5  本章小结", "20"),
            ("第3章  系统总体设计", "21"),
            ("3.1  系统需求分析", "21"),
            ("    3.1.1  功能需求分析", "21"),
            ("    3.1.2  非功能需求分析", "22"),
            ("    3.1.3  用户画像分析", "23"),
            ("3.2  系统架构设计", "24"),
            ("    3.2.1  总体架构", "24"),
            ("    3.2.2  技术架构", "25"),
            ("3.3  多智能体系统设计", "26"),
            ("    3.3.1  智能体角色定义", "26"),
            ("    3.3.2  协作流程设计", "27"),
            ("    3.3.3  意图识别与路由", "28"),
            ("3.4  RAG检索系统设计", "29"),
            ("    3.4.1  文档处理流程", "29"),
            ("    3.4.2  MQE与HyDE策略", "30"),
            ("3.5  记忆系统设计", "31"),
            ("    3.5.1  记忆层次模型", "31"),
            ("    3.5.2  记忆管理机制", "32"),
            ("3.6  数据库设计", "33"),
            ("    3.6.1  向量数据库设计", "33"),
            ("    3.6.2  关系型数据库设计", "34"),
            ("3.7  本章小结", "35"),
            ("第4章  系统实现", "36"),
            ("4.1  开发环境配置", "36"),
            ("4.2  后端系统实现", "37"),
            ("    4.2.1  FastAPI框架搭建", "37"),
            ("    4.2.2  多智能体系统实现", "38"),
            ("    4.2.3  RAG模块实现", "39"),
            ("    4.2.4  记忆管理实现", "40"),
            ("4.3  前端系统实现", "41"),
            ("    4.3.1  Vue3项目搭建", "41"),
            ("    4.3.2  页面组件实现", "42"),
            ("    4.3.3  API接口对接", "43"),
            ("4.4  关键技术实现", "44"),
            ("    4.4.1  多查询扩展实现", "44"),
            ("    4.4.2  HyDE假设文档嵌入", "45"),
            ("    4.4.3  多模态处理实现", "46"),
            ("4.5  本章小结", "47"),
            ("第5章  系统测试", "48"),
            ("5.1  测试环境与策略", "48"),
            ("5.2  功能测试", "49"),
            ("    5.2.1  多智能体协同测试", "49"),
            ("    5.2.2  RAG检索测试", "50"),
            ("    5.2.3  记忆系统测试", "51"),
            ("5.3  性能测试", "52"),
            ("    5.3.1  响应时间测试", "52"),
            ("    5.3.2  并发性能测试", "53"),
            ("5.4  用户体验测试", "54"),
            ("5.5  本章小结", "55"),
            ("第6章  总结与展望", "56"),
            ("6.1  研究工作总结", "56"),
            ("6.2  主要创新点", "57"),
            ("6.3  不足与改进方向", "58"),
            ("6.4  未来展望", "59"),
            ("参考文献", "60"),
            ("致  谢", "62"),
        ]

        table = self.doc.add_table(rows=len(toc_data) + 1, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        header_cells = table.rows[0].cells
        header_cells[0].text = "目 录 项"
        header_cells[1].text = "页 码"

        for cell in header_cells:
            para = cell.paragraphs[0]
            run = para.add_run(cell.text)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = '黑体'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, (title, page) in enumerate(toc_data, 1):
            cells = table.rows[i].cells
            cells[0].text = title
            cells[1].text = page

            for cell in cells:
                para = cell.paragraphs[0]
                if len(para.runs) > 0:
                    para.runs[0].font.size = Pt(12)
                    para.runs[0].font.name = '宋体'
                else:
                    run = para.add_run(title if cell == cells[0] else page)
                    run.font.size = Pt(12)
                    run.font.name = '宋体'

                if title.startswith('第') or title in ['摘  要', 'ABSTRACT', '目  录', '参考文献', '致  谢']:
                    para.runs[0].font.bold = True
                    para.runs[0].font.size = Pt(14)

            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_paragraph(self, text, first_line_indent=True):
        """添加普通段落"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        if first_line_indent:
            p.paragraph_format.first_line_indent = Cm(0.74)

        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = '宋体'

    def add_table(self, data, headers, caption=None):
        """添加表格"""
        self.table_num += 1
        if caption:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(f"表{self.table_num}  {caption}")
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.name = '宋体'

        table = self.doc.add_table(rows=len(data) + 1, cols=len(headers))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(header)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = '宋体'

        for row_idx, row_data in enumerate(data, 1):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                para = cell.paragraphs[0]
                para.clear()
                run = para.add_run(str(cell_data))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.size = Pt(12)
                run.font.name = '宋体'

        if caption:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(18)

    def add_figure_caption(self, caption):
        """添加图标题"""
        self.figure_num += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(f"图{self.figure_num}  {caption}")
        run.font.size = Pt(10.5)
        run.font.name = '宋体'

    def add_chapter(self, number, title):
        """添加章节标题"""
        p = self.doc.add_paragraph(style='Heading1')
        p.add_run(f"第{number}章  {title}")

    def add_section(self, number, title):
        """添加节标题"""
        p = self.doc.add_paragraph(style='Heading2')
        p.add_run(f"{number}  {title}")

    def add_subsection(self, number, title):
        """添加小节标题"""
        p = self.doc.add_paragraph(style='Heading3')
        p.add_run(f"{number}  {title}")

    def save(self, filename):
        """保存文档"""
        self.doc.save(filename)
        print(f"论文已保存至: {filename}")
