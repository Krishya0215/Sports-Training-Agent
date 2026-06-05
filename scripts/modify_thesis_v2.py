#!/usr/bin/env python3
"""
论文修改脚本 v2 —— 根据评审老师9条意见逐条修改
输入: 毕设/基于多智能体与检索增强生成的智能运动训练系统_修改版.docx
输出: 毕设/基于多智能体与检索增强生成的智能运动训练系统_修改版_v2.docx
"""

import os
import re
import shutil
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

THESIS_DIR = Path.home() / "毕设"
INPUT_FILE = THESIS_DIR / "基于多智能体与检索增强生成的智能运动训练系统_修改版.docx"
OUTPUT_FILE = THESIS_DIR / "基于多智能体与检索增强生成的智能运动训练系统_修改版_v2.docx"
DIAGRAMS_DIR = Path(__file__).parent / "diagrams_output"
DIAGRAMS_DIR.mkdir(exist_ok=True)


# ═══════════════════════════════════════════════════════════
# Helper Functions (from modify_thesis.py)
# ═══════════════════════════════════════════════════════════

def find_paragraph_index(doc, text_fragment, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if text_fragment in p.text:
            return i
    return -1


# TOC boundary: cached after first call
_toc_end_cache = {}

def get_toc_end(doc):
    """Find the last TOC paragraph index."""
    doc_id = id(doc)
    if doc_id not in _toc_end_cache:
        last_toc = 0
        for i, p in enumerate(doc.paragraphs):
            if "toc" in (p.style.name or "").lower():
                last_toc = i
        _toc_end_cache[doc_id] = last_toc
    return _toc_end_cache[doc_id]


def find_body_paragraph_index(doc, text_fragment, start=None):
    """Like find_paragraph_index but skips all TOC paragraphs."""
    body_start = get_toc_end(doc) + 1
    effective_start = max(body_start, start or 0)
    return find_paragraph_index(doc, text_fragment, start=effective_start)


def insert_paragraph_after(doc, ref_paragraph, text, style_name="Normal"):
    new_p = OxmlElement("w:p")
    ref_paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, ref_paragraph._parent)
    new_para.style = doc.styles[style_name]
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_before(doc, ref_paragraph, text, style_name="Normal"):
    new_p = OxmlElement("w:p")
    ref_paragraph._element.addprevious(new_p)
    new_para = Paragraph(new_p, ref_paragraph._parent)
    new_para.style = doc.styles[style_name]
    if text:
        new_para.add_run(text)
    return new_para


def replace_paragraph_text(paragraph, new_text):
    """Replace all text in a paragraph, clearing hyperlink runs too."""
    from lxml import etree
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    p_elem = paragraph._element

    # Collect formatting from the first run (if any) for preservation
    first_rPr = None
    first_r = p_elem.find(".//w:r", nsmap)
    if first_r is not None:
        rPr = first_r.find("w:rPr", nsmap)
        if rPr is not None:
            first_rPr = deepcopy(rPr)

    # Remove all runs (w:r) and hyperlinks (w:hyperlink) from the paragraph
    for child in list(p_elem):
        tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ""
        if tag in ("r", "hyperlink"):
            p_elem.remove(child)

    # Add a fresh run with the new text
    new_r = OxmlElement("w:r")
    if first_rPr is not None:
        new_r.insert(0, first_rPr)
    new_t = OxmlElement("w:t")
    new_t.text = new_text
    new_t.set(qn("xml:space"), "preserve")
    new_r.append(new_t)
    p_elem.append(new_r)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)


def insert_image_after(doc, ref_paragraph, image_path, width_cm=14):
    """Insert an image centered in a new paragraph after ref_paragraph."""
    new_p = OxmlElement("w:p")
    ref_paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, ref_paragraph._parent)
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return new_para


# ═══════════════════════════════════════════════════════════
# Diagram Generation
# ═══════════════════════════════════════════════════════════

def generate_use_case_diagram():
    """意见6: 生成用例图 PNG"""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.patches import FancyBboxPatch

    plt.rcParams["font.family"] = ["Hiragino Sans GB", "Songti SC", "Heiti TC", "STHeiti", "Arial Unicode MS", "sans-serif"]
    plt.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(1, 1, figsize=(12, 10))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 12)
    ax.set_aspect("equal")
    ax.axis("off")

    # System boundary
    rect = FancyBboxPatch((2.5, 0.5), 7, 11, boxstyle="round,pad=0.3",
                           edgecolor="black", facecolor="#f9f9f9", linewidth=1.5)
    ax.add_patch(rect)
    ax.text(6, 11.2, "智能运动训练系统", ha="center", fontsize=13, fontweight="bold")

    # Actor: User (left)
    def draw_actor(ax, x, y, label):
        ax.plot(x, y + 0.3, "o", color="black", markersize=8)
        ax.plot([x, x], [y - 0.1, y + 0.2], color="black", linewidth=1.5)
        ax.plot([x - 0.3, x + 0.3], [y + 0.05, y + 0.05], color="black", linewidth=1.5)
        ax.plot([x - 0.2, x], [y - 0.4, y - 0.1], color="black", linewidth=1.5)
        ax.plot([x + 0.2, x], [y - 0.4, y - 0.1], color="black", linewidth=1.5)
        ax.text(x, y - 0.65, label, ha="center", fontsize=10)

    draw_actor(ax, 1.2, 6, "用户")
    draw_actor(ax, 10.8, 7, "管理员")

    # Use cases
    user_cases = [
        (6, 10.3, "注册与登录"),
        (6, 9.2, "个人资料管理"),
        (6, 8.1, "AI 智能问答"),
        (6, 7.0, "训练计划管理"),
        (6, 5.9, "训练记录管理"),
        (6, 4.8, "饮食记录管理"),
        (6, 3.7, "体重记录管理"),
        (6, 2.6, "数据统计分析"),
        (6, 1.5, "个人记忆查看"),
    ]
    admin_cases = [
        (6, 10.3, "注册与登录"),
        (8.2, 6.5, "知识库文档管理"),
    ]

    for cx, cy, label in user_cases:
        ellipse = mpatches.Ellipse((cx, cy), 3.0, 0.7, edgecolor="black",
                                    facecolor="white", linewidth=1)
        ax.add_patch(ellipse)
        ax.text(cx, cy, label, ha="center", va="center", fontsize=9)
        ax.plot([1.5, cx - 1.5], [6, cy], color="gray", linewidth=0.6)

    # Admin use case
    ellipse = mpatches.Ellipse((8.2, 6.5), 3.0, 0.7, edgecolor="black",
                                facecolor="white", linewidth=1)
    ax.add_patch(ellipse)
    ax.text(8.2, 6.5, "知识库文档管理", ha="center", va="center", fontsize=9)
    ax.plot([10.5, 8.2 + 1.5], [7, 6.5], color="gray", linewidth=0.6)
    # Admin also has login
    ax.plot([10.5, 6 + 1.5], [7, 10.3], color="gray", linewidth=0.6)

    out = DIAGRAMS_DIR / "use_case_diagram.png"
    fig.savefig(out, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"  [OK] 用例图: {out}")
    return out


def generate_er_diagram():
    """意见8: 生成陈氏ER图 (实体=矩形, 属性=椭圆, 关系=菱形)"""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np

    plt.rcParams["font.family"] = ["Hiragino Sans GB", "Songti SC", "Heiti TC",
                                    "STHeiti", "Arial Unicode MS", "sans-serif"]
    plt.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(1, 1, figsize=(26, 20))
    ax.set_xlim(0, 26)
    ax.set_ylim(0, 20)
    ax.set_aspect("equal")
    ax.axis("off")

    # ── Colors ──
    ENT_C = "#2B6CA3"
    ENT_T = "white"
    ATTR_C = "#6B96B2"
    ATTR_T = "white"
    PK_C = "#C9983A"
    REL_C = "#E8A83E"
    REL_T = "white"
    LINE_C = "#555555"

    # ── Drawing helpers ──
    def draw_entity(x, y, name, w=2.4, h=0.8):
        r = mpatches.FancyBboxPatch((x - w/2, y - h/2), w, h,
                boxstyle="round,pad=0.05", fc=ENT_C, ec="white", lw=2)
        ax.add_patch(r)
        ax.text(x, y, name, ha="center", va="center", fontsize=11,
                fontweight="bold", color=ENT_T)

    def draw_attr(x, y, name, is_pk=False):
        c = PK_C if is_pk else ATTR_C
        e = mpatches.Ellipse((x, y), 1.7, 0.58, fc=c, ec="white", lw=0.8)
        ax.add_patch(e)
        fs = 7.5 if len(name) <= 6 else 6.8
        ax.text(x, y, name, ha="center", va="center", fontsize=fs, color=ATTR_T)

    def draw_rel(x, y, name):
        pts = np.array([[x, y+0.52], [x+0.72, y], [x, y-0.52], [x-0.72, y]])
        diamond = mpatches.Polygon(pts, closed=True, fc=REL_C, ec="white", lw=1.5)
        ax.add_patch(diamond)
        ax.text(x, y, name, ha="center", va="center", fontsize=9,
                fontweight="bold", color=REL_T)

    def line(x1, y1, x2, y2):
        ax.plot([x1, x2], [y1, y2], color=LINE_C, lw=0.9, zorder=0)

    def card(x, y, txt):
        ax.text(x, y, txt, ha="center", va="center", fontsize=7.5,
                color="#333", fontweight="bold")

    # ── Title ──
    ax.text(13, 19.5, "智能运动训练系统ER图", fontsize=16, fontweight="bold", ha="center")

    # ══════════════════════════════════════════════════════
    # Entity positions
    # ══════════════════════════════════════════════════════
    E = {
        "用户":     (13, 14.5),
        "用户档案": (5, 17),
        "训练计划": (3.5, 11),
        "训练记录": (3.5, 5.5),
        "情景记忆": (22.5, 17),
        "语义记忆": (22.5, 11),
        "工作会话": (22.5, 5.5),
        "对话记录": (13, 2),
    }
    for name, (ex, ey) in E.items():
        draw_entity(ex, ey, name)

    # ══════════════════════════════════════════════════════
    # Attributes for each entity  (x, y, label, is_pk)
    # ══════════════════════════════════════════════════════
    attrs = {
        "用户": [
            (13, 16.8, "用户ID", True),
            (10.8, 16.1, "用户名", False),
            (15.2, 16.1, "邮箱", False),
            (10.4, 15.2, "密码", False),
            (15.6, 15.2, "角色", False),
            (13, 13.0, "头像", False),
        ],
        "用户档案": [
            (3.0, 18.7, "档案ID", True),
            (5.0, 18.7, "训练目标", False),
            (7.0, 18.5, "训练水平", False),
            (2.2, 17.8, "伤病状态", False),
            (7.2, 17.6, "偏好方式", False),
            (3.5, 15.8, "每周天数", False),
            (6.5, 16.0, "伤病详情", False),
        ],
        "训练计划": [
            (1.2, 12.4, "计划ID", True),
            (3.5, 12.6, "标题", False),
            (5.6, 12.0, "目标", False),
            (1.2, 11.0, "状态", False),
            (1.2, 9.8, "版本", False),
            (5.4, 10.5, "内容", False),
        ],
        "训练记录": [
            (1.2, 6.8, "记录ID", True),
            (5.6, 6.8, "日期", False),
            (1.0, 5.5, "训练类型", False),
            (1.2, 4.3, "疲劳等级", False),
            (5.6, 4.4, "疼痛等级", False),
            (3.5, 3.8, "完成状态", False),
        ],
        "情景记忆": [
            (20.3, 18.7, "事件ID", True),
            (22.5, 18.7, "事件类型", False),
            (24.5, 18.5, "事件时间", False),
            (20.0, 17.8, "事件摘要", False),
            (24.8, 17.5, "重要性评分", False),
        ],
        "语义记忆": [
            (20.0, 12.5, "事实ID", True),
            (24.8, 12.3, "事实类别", False),
            (20.2, 10.0, "事实键", False),
            (22.5, 9.6, "事实值", False),
            (24.8, 10.2, "置信度", False),
        ],
        "工作会话": [
            (20.0, 6.8, "会话ID", True),
            (24.8, 6.6, "对话标识", False),
            (20.0, 4.5, "来源", False),
            (22.5, 4.0, "状态", False),
            (24.8, 4.8, "最大轮数", False),
        ],
        "对话记录": [
            (10.5, 2.0, "消息ID", True),
            (15.5, 2.0, "会话标识", False),
            (10.5, 0.8, "问题", False),
            (13, 0.5, "回答", False),
            (15.5, 0.8, "模式", False),
        ],
    }

    for ent_name, alist in attrs.items():
        ex, ey = E[ent_name]
        for ax_, ay_, lbl, pk in alist:
            draw_attr(ax_, ay_, lbl, pk)
            line(ex, ey, ax_, ay_)

    # ══════════════════════════════════════════════════════
    # Relationships + connecting lines + cardinality
    # ══════════════════════════════════════════════════════
    # 拥有: 用户 <-> 用户档案 (1:1)
    rx, ry = 8.5, 16.2
    draw_rel(rx, ry, "拥有")
    line(E["用户"][0], E["用户"][1], rx, ry)
    line(E["用户档案"][0], E["用户档案"][1], rx, ry)
    card(10.8, 16.6, "1")
    card(6.7, 16.7, "1")

    # 制定: 用户 -> 训练计划 (1:N)
    rx, ry = 8, 12.8
    draw_rel(rx, ry, "制定")
    line(E["用户"][0], E["用户"][1], rx, ry)
    line(E["训练计划"][0], E["训练计划"][1], rx, ry)
    card(10.4, 13.9, "1")
    card(5.8, 12.0, "N")

    # 包含: 训练计划 -> 训练记录 (1:N)
    rx, ry = 3.5, 8.3
    draw_rel(rx, ry, "包含")
    line(E["训练计划"][0], E["训练计划"][1], rx, ry)
    line(E["训练记录"][0], E["训练记录"][1], rx, ry)
    card(3.9, 9.8, "1")
    card(3.9, 6.8, "N")

    # 产生: 用户 -> 情景记忆 (1:N)
    rx, ry = 18, 16.2
    draw_rel(rx, ry, "产生")
    line(E["用户"][0], E["用户"][1], rx, ry)
    line(E["情景记忆"][0], E["情景记忆"][1], rx, ry)
    card(15.3, 15.7, "1")
    card(20.5, 16.8, "N")

    # 归纳: 情景记忆 -> 语义记忆 (1:N)
    rx, ry = 22.5, 14
    draw_rel(rx, ry, "归纳")
    line(E["情景记忆"][0], E["情景记忆"][1], rx, ry)
    line(E["语义记忆"][0], E["语义记忆"][1], rx, ry)
    card(22.9, 15.7, "1")
    card(22.9, 12.3, "N")

    # 对话: 用户 -> 对话记录 (1:N)
    rx, ry = 10, 8.5
    draw_rel(rx, ry, "对话")
    line(E["用户"][0], E["用户"][1], rx, ry)
    line(E["对话记录"][0], E["对话记录"][1], rx, ry)
    card(11.4, 11.5, "1")
    card(11.4, 5.2, "N")

    # 开启: 用户 -> 工作会话 (1:N)
    rx, ry = 18, 9.5
    draw_rel(rx, ry, "开启")
    line(E["用户"][0], E["用户"][1], rx, ry)
    line(E["工作会话"][0], E["工作会话"][1], rx, ry)
    card(15.6, 12.0, "1")
    card(20.2, 7.5, "N")

    # 记录: 用户 -> 训练记录 (1:N)  (user also directly owns records)
    rx, ry = 8, 9.5
    draw_rel(rx, ry, "记录")
    line(E["用户"][0], E["用户"][1], rx, ry)
    line(E["训练记录"][0], E["训练记录"][1], rx, ry)
    card(10.3, 12.0, "1")
    card(5.5, 7.5, "N")

    out = DIAGRAMS_DIR / "er_diagram.png"
    fig.savefig(out, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"  [OK] ER图: {out}")
    return out


def _generate_er_matplotlib():
    """Unused — Chen-style ER is now in generate_er_diagram()."""
    pass


def generate_architecture_diagram():
    """意见7: 生成正确的系统架构图"""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

    plt.rcParams["font.family"] = ["Hiragino Sans GB", "Songti SC", "Heiti TC", "STHeiti", "Arial Unicode MS", "sans-serif"]
    plt.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis("off")

    def draw_layer(ax, x, y, w, h, title, items, color):
        box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                              edgecolor="#333", facecolor=color, linewidth=1.5)
        ax.add_patch(box)
        ax.text(x + w / 2, y + h - 0.25, title, ha="center", va="top",
                fontsize=11, fontweight="bold")
        for i, item in enumerate(items):
            ix = x + 0.5 + (i % 3) * (w - 1) / 3
            iy = y + h - 0.7 - (i // 3) * 0.5
            item_box = FancyBboxPatch((ix - 0.1, iy - 0.15), (w - 1) / 3 - 0.2, 0.35,
                                       boxstyle="round,pad=0.05", edgecolor="#666",
                                       facecolor="white", linewidth=0.8)
            ax.add_patch(item_box)
            ax.text(ix + (w - 1) / 6 - 0.2, iy + 0.02, item, ha="center", va="center", fontsize=8)

    # Layer 1: Frontend
    draw_layer(ax, 1, 7.2, 12, 1.3, "前端展示层",
               ["Vue3 + Vite", "Axios HTTP", "Markdown渲染"], "#DCEDC8")

    # Layer 2: Backend (CENTRAL)
    draw_layer(ax, 1, 4.2, 12, 2.5, "后端服务层（FastAPI）",
               ["RAG 检索模块", "长期记忆模块", "多智能体模块",
                "用户认证模块", "训练计划模块", "数据记录模块"], "#BBDEFB")

    # Layer 3a: Data Storage (left)
    draw_layer(ax, 1, 1, 5.5, 2.5, "数据存储层",
               ["SQLite 关系数据库", "ChromaDB 向量数据库"], "#FFF9C4")

    # Layer 3b: Model Services (right)
    draw_layer(ax, 7, 1, 6, 2.5, "模型服务层",
               ["LLM API (通义千问)", "Embedding API"], "#F8BBD0")

    # Arrows: Frontend <-> Backend
    ax.annotate("", xy=(7, 7.2), xytext=(7, 6.7),
                arrowprops=dict(arrowstyle="<->", color="#333", lw=2))
    ax.text(7.8, 6.9, "REST API", fontsize=8, color="#555")

    # Arrows: Backend -> Data Storage
    ax.annotate("", xy=(3.75, 3.5), xytext=(3.75, 4.2),
                arrowprops=dict(arrowstyle="<->", color="#333", lw=2))
    ax.text(4.1, 3.75, "读写数据", fontsize=8, color="#555")

    # Arrows: Backend -> Model Services
    ax.annotate("", xy=(10, 3.5), xytext=(10, 4.2),
                arrowprops=dict(arrowstyle="<->", color="#333", lw=2))
    ax.text(10.4, 3.75, "模型调用", fontsize=8, color="#555")

    # NO arrow between Data Storage and Model Services!
    # Add a note
    ax.text(7, 0.4, "注：数据存储层与模型服务层之间无直接交互，所有数据流均经由后端服务层调度",
            ha="center", fontsize=8, style="italic", color="#888")

    out = DIAGRAMS_DIR / "architecture_diagram.png"
    fig.savefig(out, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"  [OK] 架构图: {out}")
    return out


# ═══════════════════════════════════════════════════════════
# Phase 1: Text Modifications
# ═══════════════════════════════════════════════════════════

def task_1_1_toc_cleanup(doc):
    """意见2: 目录去掉英文缩写标注"""
    print("\n[Task 1.1] 意见2: 目录去掉英文缩写标注")
    count = 0
    for p in doc.paragraphs:
        style = p.style.name.lower()
        if "toc" in style or p.style.name in ("2级标题", "3级标题", "1级标题"):
            old = p.text
            # Remove parenthetical English abbreviations like （RAG）, (MAS), （HyDE）
            new = re.sub(r'[（(]\s*[A-Za-z][A-Za-z\s\-]*\s*[)）]', '', old)
            if new != old:
                replace_paragraph_text(p, new)
                count += 1
                print(f"  修改: '{old.strip()[:60]}' -> '{new.strip()[:60]}'")
    print(f"  共修改 {count} 处 TOC/标题")


def task_1_2_distinguish_retain_utilize(doc):
    """意见1: 区分'保留'与'利用'"""
    print("\n[Task 1.2] 意见1: 区分'保留'与'利用'")

    # Abstract (before TOC, use find_paragraph_index)
    idx = find_paragraph_index(doc, "系统无法跨会话保留用户状态")
    if idx >= 0:
        p = doc.paragraphs[idx]
        old = p.text
        new = old.replace(
            "系统无法跨会话保留用户状态，难以实现真正的个性化管理",
            "现有系统面临两层递进不足：一是无法跨会话持久化用户的训练偏好、伤病约束与历史状态；"
            "二是即使具备基本存储能力，也缺乏结构化地利用这些长期状态进行个性化检索与生成的机制"
        )
        if new != old:
            replace_paragraph_text(p, new)
            print(f"  摘要已修改 (段落 {idx})")
    else:
        print("  [WARN] 未找到摘要中的目标文本")

    # Tech section (also in abstract area, before TOC)
    idx2 = find_paragraph_index(doc, "针对系统无法持续利用用户长期状态的问题")
    if idx2 >= 0:
        p = doc.paragraphs[idx2]
        old = p.text
        new = old.replace(
            "针对系统无法持续利用用户长期状态的问题",
            "针对用户长期状态难以被结构化建模与主动利用的问题"
        )
        if new != old:
            replace_paragraph_text(p, new)
            print(f"  技术方案已修改 (段落 {idx2})")
    else:
        print("  [WARN] 未找到技术方案中的目标文本")


def task_1_3_rewrite_problem2(doc):
    """意见4: 问题二描述更客观"""
    print("\n[Task 1.3] 意见4: 重写问题二（更客观地描述LLM记忆现状）")

    idx = find_body_paragraph_index(doc, "当前大多数大语言模型在交互层面仍以")
    if idx < 0:
        idx = find_body_paragraph_index(doc, "缺乏长期用户信息记忆")
    if idx < 0:
        print("  [WARN] 未找到问题二段落，尝试模糊匹配")
        idx = find_body_paragraph_index(doc, "问题二")
    if idx < 0:
        print("  [ERROR] 无法定位问题二段落，跳过")
        return

    new_text = (
        "问题二：用户长期状态缺乏结构化建模，难以驱动个性化训练决策。"
        "近年来，部分商用大语言模型（如ChatGPT的Memory功能、Claude的Projects机制）"
        "已具备基础的跨会话记忆能力，能在对话间保留用户提及的偏好信息。"
        "然而，这些通用记忆机制以非结构化的自由文本形式存储信息，存在三方面局限："
        "其一，缺乏领域专属的结构化表示，无法以量化方式建模运动训练中的伤病约束、"
        "疲劳评分、训练完成率等关键指标；"
        "其二，不具备重要性评估与时间衰减机制，难以区分核心信息与过时信息；"
        "其三，通用记忆无法与检索增强生成系统深度联动，"
        "不能将用户的结构化状态注入检索查询以实现个性化知识召回。"
        "运动训练是一个持续数周甚至数月的动态过程，"
        "训练方案的合理性不仅取决于当前提问内容，"
        "还与用户的长期目标、伤病历史、恢复状态及近期训练负荷密切相关。"
        "在缺乏领域专属记忆架构的情况下，系统难以形成稳定的用户画像，"
        "生成的建议也缺乏对个体差异的持续追踪与适应。"
        "如何构建能够跨会话积累并主动利用结构化用户状态的记忆机制，"
        "是实现个性化训练的重要前提。"
    )
    replace_paragraph_text(doc.paragraphs[idx], new_text)
    print(f"  问题二已重写 (段落 {idx})")


def task_1_4_rewrite_section_1_5(doc):
    """意见5: 重写1.5节，消除LLM生成痕迹"""
    print("\n[Task 1.4] 意见5: 重写1.5节")

    # Find section 1.5 heading
    sec_start = find_body_paragraph_index(doc, "技术方案与关键挑战分析")
    if sec_start < 0:
        print("  [ERROR] 未找到1.5节标题")
        return

    # Find the next section (1.6 论文组织结构)
    sec_end = find_body_paragraph_index(doc, "论文组织结构", start=sec_start + 1)
    if sec_end < 0:
        print("  [ERROR] 未找到1.6节标题")
        return

    # Delete all paragraphs between 1.5 heading and 1.6 heading (exclusive)
    # We need to work backwards to avoid index shifts
    paras_to_delete = []
    for i in range(sec_start + 1, sec_end):
        if doc.paragraphs[i].text.strip():
            paras_to_delete.append(i)

    for i in sorted(paras_to_delete, reverse=True):
        delete_paragraph(doc.paragraphs[i])

    print(f"  已删除 {len(paras_to_delete)} 个段落 ({sec_start+1} ~ {sec_end-1})")

    # Insert new content after the section heading
    ref = doc.paragraphs[sec_start]  # The heading paragraph

    new_paragraphs = [
        # 1.5 引导段
        ("从识别问题到提出可行方案，中间还需要进一步回答两个关键问题：这些问题在技术层面究竟对应什么样的求解目标；所选技术在运动训练这一具体场景中落地时，会遇到哪些实际困难。本节将沿着这两条线索，逐一分析三项技术方案的设计动机与面临的核心挑战。", "Normal"),

        ("通过对现有智能运动训练系统与通用检索增强生成系统的对比分析，可以发现运动训练场景中的核心矛盾并非仅仅是“大模型知识不足”这一单一问题。实际上，系统需要同时弥合三层错位：用户训练语言与专业知识文档之间的语义鸿沟、用户长期训练状态在跨周期演化中的建模缺失、以及运动训练问题天然具有的多目标多角色耦合特征。无论是直接迁移通用检索增强生成方案，还是单纯堆叠更大规模的模型，都难以同时稳定满足训练建议的科学性、个性化连续性与多专业覆盖。基于此判断，本文围绕训练知识检索、长期用户状态建模与多角色协同决策三个方向展开技术方案设计。", "Normal"),

        # 1.5.1
        ("面向运动训练场景的语义增强RAG", "3级标题"),

        ("问题一的本质，是用户自然语言表达与专业运动知识之间的语义错位。传统检索增强生成系统默认用户查询与知识库文档处于相同语义空间，只需通过向量相似度即可完成有效召回。在运动训练场景中，这一假设并不成立。普通用户在提问时倾向于使用口语化表达（如“练完腿第二天特别酸怎么办”），而知识库中的专业文档使用的是“延迟性肌肉酸痛”“恢复周期”等术语。两者之间存在的语义差距，导致传统向量检索容易出现召回偏移与知识遗漏。", "Normal"),

        ("除语义差距外，运动训练问题还具有多维耦合特征。以“减脂期间深蹲后膝盖疼怎么办”为例，该问题同时涉及减脂策略、力量训练负荷、深蹲动作技术、膝关节生物力学与损伤预防等多个维度，传统单次检索往往只能覆盖其中部分语义子空间。", "Normal"),

        ("针对上述两类挑战，本文设计面向运动训练的语义增强检索增强生成策略。在检索侧，引入假设文档嵌入机制，由大语言模型根据用户问题生成一段假设性专业回答，将其映射到专业知识空间后再执行检索，从而缩小用户表达与专业文档之间的分布差异。在此基础上，采用多查询扩展机制将复杂训练问题拆解为多个语义子查询并行检索，提升多维训练问题下的知识覆盖能力。在生成侧，通过提示词约束机制要求模型优先依据检索文档作答，并在后处理阶段增加引用一致性校验。该策略围绕运动训练语义鸿沟这一场景问题展开针对性优化，而非简单叠加多个检索技巧。", "Normal"),

        # 1.5.2
        ("基于长期记忆的用户状态建模", "3级标题"),

        ("运动训练并非一次性问答，而是持续数周甚至数月的动态过程。用户当前训练方案的合理性，不仅取决于当前提问内容，还与长期目标、训练水平、伤病历史、恢复状态及近期训练负荷密切相关。这意味着，即使检索增强生成系统能够准确召回专业知识，若忽略用户个体状态，也可能给出不适用于当前用户的建议——对于同样的“今天怎么练腿”问题，存在膝关节旧伤的用户与竞技力量训练者应当得到完全不同的回答。", "Normal"),

        ("在实际开发中，我们发现传统检索增强生成的检索过程仅依赖当前查询，无法感知用户长期状态。为此，本文提出记忆感知检索机制：在检索增强生成检索前，从长期记忆中提取用户状态摘要，与当前问题共同拼接为增强查询，实现“Query + User State”的联合检索。具体而言，本文设计由工作记忆、情景记忆与语义记忆构成的三层记忆体系：工作记忆维护当前会话上下文；情景记忆保存用户历史训练事件与身体反馈记录；语义记忆抽象存储用户的长期稳定特征（训练目标、运动偏好、伤病历史等）。为避免记忆无限膨胀，系统引入四维重要性评分（核心目标相关度、伤病风险、重复频次、决策影响度）控制写入，并通过基于时间衰减与访问频率的遗忘机制对低强度记忆进行归档。", "Normal"),

        # 1.5.3
        ("基于多智能体协同的决策分工", "3级标题"),

        ("运动训练本质上是一个多角色协同的过程。现实中的专业训练团队通常包含负责周期规划与负荷控制的训练规划师、负责动作模式与细节优化的技术教练、以及负责伤病风险评估与恢复策略的运动康复师。这三个角色在知识结构与决策目标上存在显著差异，甚至可能产生冲突——例如，训练规划师为提高力量水平安排的大负荷深蹲，可能被康复师因膝关节风险而否决。如果将所有职责交由单一模型处理，模型在面对复杂问题时往往顾此失彼。", "Normal"),

        ("本文采用基于LangGraph的多智能体协同架构来应对这一挑战。系统将训练规划教练、技术指导教练与运动康复教练设计为独立角色智能体，每个智能体拥有专属的提示词模板、推理参数与领域约束。在具体流程中，系统首先通过多标签意图识别判断用户问题涉及的专业维度；调度模块据此动态激活对应智能体；各智能体基于自身职责分别执行检索与推理；最后由综合输出器对多角色结果进行统一整合与冲突消解。当用户提出“减脂期间跑步膝盖疼怎么办”时，系统会同时激活训练规划智能体分析减脂训练负荷与康复智能体评估膝关节风险，从而避免单角色视角导致的建议偏差。为控制多智能体调用成本，系统采用条件激活策略：仅当问题涉及多个专业维度时才进入多智能体模式，简单问题由单智能体快速响应。", "Normal"),

        # 1.5.4
        ("三者协同的技术闭环", "3级标题"),

        ("语义增强检索增强生成、长期记忆与多智能体并非彼此独立的模块。在系统运行中，长期记忆参与检索过程，使检索结果与用户个体状态相匹配；多智能体内部共享统一的用户状态对象，各教练在推理时均可获取该用户的伤病约束与训练偏好；检索结果又为各教练提供了经过知识验证的专业依据。三者分别对应“知识失配”“状态失配”“角色失配”三类核心问题，共同构成系统的技术闭环。", "Normal"),
    ]

    # Insert from bottom to top (reversed) to maintain correct order
    for text, style in reversed(new_paragraphs):
        ref = insert_paragraph_after(doc, doc.paragraphs[sec_start], text, style)

    print(f"  已插入 {len(new_paragraphs)} 个新段落")


# ═══════════════════════════════════════════════════════════
# Phase 2: Content Insertions & Diagrams
# ═══════════════════════════════════════════════════════════

def task_2_1_add_rag_pipeline_description(doc):
    """意见9: 在3.3.1节补充RAG数据处理流水线"""
    print("\n[Task 2.1] 意见9: 补充RAG数据处理流水线描述")

    # Find insertion point: before "查询增强阶段" or after RAG section heading
    idx = find_body_paragraph_index(doc, "针对知识失配的RAG检索模块设计")
    if idx < 0:
        idx = find_body_paragraph_index(doc, "RAG检索模块设计")
    if idx < 0:
        print("  [ERROR] 未找到RAG模块设计节")
        return

    # Find the first mention of "查询增强" to insert before it
    query_enhance_idx = find_body_paragraph_index(doc, "查询增强阶段", start=idx)
    if query_enhance_idx < 0:
        query_enhance_idx = find_body_paragraph_index(doc, "（1）查询增强", start=idx)
    if query_enhance_idx < 0:
        # Insert after the section intro paragraph instead
        query_enhance_idx = idx + 2

    insert_ref = doc.paragraphs[query_enhance_idx]

    paragraphs = [
        ("在检索增强生成链路中，知识库的数据处理质量直接决定了下游检索的有效性。"
         "本文设计了一条完整的文档处理流水线，覆盖从原始文档到可检索向量的全流程。", "Normal"),

        ("在文档解析阶段，系统支持PDF、Markdown和纯文本三种知识库文件格式。"
         "其中，PDF文档的解析采用三级降级策略：首先尝试使用MarkItDown将PDF转换为结构化Markdown，"
         "保留标题层级与段落结构；若MarkItDown输出为空（常见于扫描版PDF），"
         "则回退至PyPDF逐页提取文字层；若文字层仍为空，则调用Tesseract OCR对页面图像进行光学字符识别。"
         "该策略保证了对文字型PDF、混合排版PDF与纯扫描PDF的全面覆盖。", "Normal"),

        ("在文档分块阶段，系统采用两级分割策略。"
         "第一级为基于Markdown标题结构的语义分割，沿一级标题、二级标题和三级标题进行切分，"
         "保证每个分块在语义上具有内聚性。"
         "第二级为递归字符分割，对超过800字符的块进一步切分，重叠窗口设置为100字符以避免上下文断裂。"
         "分割器的分隔符优先级依次为段落换行、句号、分号和逗号，适配中文文档的排版习惯。", "Normal"),

        ("在向量化与入库阶段，系统调用text-embedding-v4嵌入模型将文本块映射为稠密向量，"
         "存入ChromaDB持久化向量数据库。每条向量记录附带来源文件路径、文件名、"
         "解析方法（markitdown/pypdf/ocr）和内容类型等元数据，用于后续检索阶段的来源追溯与过滤。"
         "为避免知识库更新时重复处理已有文档，"
         "系统在入库前计算文件的MD5指纹并与已有记录比对，仅对新增或变更的文件执行解析与入库。", "Normal"),
    ]

    for text, style in reversed(paragraphs):
        insert_paragraph_before(doc, insert_ref, text, style)

    print(f"  已在RAG模块设计节插入 {len(paragraphs)} 段数据处理描述")


def task_2_2_add_multiagent_detail(doc):
    """意见9: 在3.3.3节补充多智能体构建细节"""
    print("\n[Task 2.2] 意见9: 补充多智能体构建细节")

    # Find insertion point: after "教练智能体角色设计" or the multi-agent section
    idx = find_body_paragraph_index(doc, "针对角色失配的多智能体模块设计")
    if idx < 0:
        idx = find_body_paragraph_index(doc, "多智能体模块设计")
    if idx < 0:
        print("  [ERROR] 未找到多智能体模块设计节")
        return

    # Find "教练智能体角色设计" or insert after the first descriptive paragraph
    coach_design_idx = find_body_paragraph_index(doc, "教练智能体角色设计", start=idx)
    if coach_design_idx < 0:
        coach_design_idx = idx + 2

    insert_ref = doc.paragraphs[coach_design_idx]

    paragraphs = [
        ("在教练智能体的工程实现层面，系统设计了统一的教练智能体基类CoachAgent。"
         "该基类封装了智能体的名称、角色描述、提示词模板与推理链，"
         "并为每个教练提供独立的推理参数配置（包括生成温度、最大输出长度、请求超时时间与重试次数），"
         "使不同教练能够根据各自的专业特性采用差异化的推理策略。"
         "每个教练智能体在处理请求时，会将用户输入、用户档案、检索到的参考文档三者注入其专属提示词模板，"
         "经由大语言模型推理生成该角色视角下的专业建议。", "Normal"),

        ("三个教练智能体的专业化配置体现了角色分工的设计意图。"
         "训练规划教练的生成温度设为0.35、最大输出长度设为1400 tokens，"
         "以确保其输出的训练计划结构完整且内容详尽；"
         "其提示词中包含严格的触发约束——仅当用户明确发出“帮我制定计划”“给我一份方案”等包含"
         "动作指令与计划词组合的请求时才输出训练计划，避免在一般咨询场景下过度生成。"
         "技术指导教练的生成温度设为0.2、最大输出长度设为420 tokens，"
         "以输出精简、确定性高的动作修正建议。"
         "运动康复教练同样采用低温度、短输出的配置，"
         "且在依赖关系上标注为依赖训练规划教练与技术指导教练的输出，"
         "以便在评估损伤风险时参考前序教练给出的训练负荷与动作方案。", "Normal"),

        ("在意图识别与智能体调度方面，系统设计了基于关键词与用户档案联合判断的多标签路由机制。"
         "当用户请求进入系统后，路由模块首先扫描输入文本中的领域关键词"
         "（如“动作”“姿势”“技术”对应技术维度，“恢复”“疼痛”“损伤”对应康复维度），"
         "同时检查用户档案中是否存在伤病记录。"
         "路由模块输出有序的意图信号集合，并据此决定激活哪些教练智能体。"
         "只有当用户输入同时包含动作指令词（如“帮我制定”“给我”）和计划目标词（如“训练计划”“健身方案”）时，"
         "系统才会调度训练规划教练；若用户档案中存在伤病信息，"
         "康复教练将自动追加到激活列表中，即使用户未显式提及康复需求。"
         "这一设计使系统在多数场景下仅激活1至2个教练快速响应，"
         "而在复杂多维问题中激活全部教练进行协同推理。", "Normal"),

        ("在工作流编排层面，系统基于LangGraph构建包含五个节点的有向状态图。"
         "检索知识节点负责调用高级检索器获取相关文档；"
         "构建执行计划节点根据意图识别结果选择参与教练并安排执行批次；"
         "执行教练节点通过线程池并发调用被选中的教练，各教练独立推理后将结果写入共享状态对象；"
         "综合响应节点将多个教练的输出注入综合提示词模板，由大语言模型整合为一份连贯的最终回答；"
         "更新记忆节点将本轮交互写入情景记忆，完成记忆闭环。"
         "当选中的教练仅有一个时，系统跳过并发步骤直接串行执行，以减少不必要的线程开销。", "Normal"),
    ]

    for text, style in reversed(paragraphs):
        ref = insert_paragraph_after(doc, insert_ref, text, style)

    print(f"  已在多智能体模块设计节插入 {len(paragraphs)} 段构建细节描述")


def task_2_3_insert_use_case_diagram(doc, image_path):
    """意见6: 在3.1.1前插入用例图"""
    print("\n[Task 2.3] 意见6: 插入用例图")

    idx = find_body_paragraph_index(doc, "功能性需求分析")
    if idx < 0:
        print("  [ERROR] 未找到3.1.1节标题")
        return

    ref = doc.paragraphs[idx]
    # Insert caption after heading, then image before caption
    caption = insert_paragraph_after(doc, ref, "图3-0 系统用例图", "Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.size = Pt(9)

    img_para = insert_paragraph_after(doc, ref, "", "Normal")
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Cm(13))

    desc = insert_paragraph_after(doc, ref,
        "根据第三章各节描述的系统功能，本文首先从用户与管理员两类参与者的角度梳理系统用例，"
        "用例图如下所示。用户是系统的主要参与者，可进行注册登录、个人资料管理、AI智能问答、"
        "训练计划管理、训练记录管理、饮食记录管理、体重记录管理、数据统计分析和个人记忆查看等操作。"
        "管理员除具备用户的基本功能外，还可进行知识库文档管理操作。", "Normal")

    print(f"  已在段落 {idx} 后插入用例图")


def task_2_4_insert_er_diagram(doc, image_path):
    """意见8: 在3.4.2节插入ER图"""
    print("\n[Task 2.4] 意见8: 插入ER图")

    # Find "数据表设计" or "系统共设计" in 3.4.2
    idx = find_body_paragraph_index(doc, "系统共设计")
    if idx < 0:
        idx = find_body_paragraph_index(doc, "数据表设计")
    if idx < 0:
        idx = find_body_paragraph_index(doc, "关系数据库")
    if idx < 0:
        print("  [ERROR] 未找到数据表设计段落")
        return

    ref = doc.paragraphs[idx]

    caption = insert_paragraph_after(doc, ref, "图3-x 系统数据库ER图", "Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.size = Pt(9)

    img_para = insert_paragraph_after(doc, ref, "", "Normal")
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Cm(14))

    desc = insert_paragraph_after(doc, ref,
        "为清晰展示各数据表之间的关联关系，下图给出了系统数据库的实体关系图。"
        "图中以users表为核心，通过外键约束与其他13张表建立关联。"
        "其中，user_profiles与users为一对一关系，存储用户长期画像信息；"
        "training_plans、training_records、chat_messages、daily_records和weight_records"
        "均与users为一对多关系，记录用户的训练、对话与健康数据。"
        "在记忆表族中，memory_episodic_events通过user_id关联用户，"
        "memory_semantic_facts通过source_event_id与情景事件表建立引用关系，"
        "memory_working_sessions和memory_working_messages则构成会话级的工作记忆层级结构。", "Normal")

    print(f"  已在段落 {idx} 后插入ER图")


def task_2_5_replace_architecture_diagram(doc, image_path):
    """意见7: 替换架构图（图3-1）"""
    print("\n[Task 2.5] 意见7: 替换架构图")

    # Find the paragraph containing "图3-1" caption or the architecture description
    idx = find_body_paragraph_index(doc, "图3-1")
    if idx < 0:
        idx = find_body_paragraph_index(doc, "整体架构")
    if idx < 0:
        print("  [ERROR] 未找到架构图位置")
        return

    # Find the architecture description paragraph to insert the new diagram
    arch_desc_idx = find_body_paragraph_index(doc, "自上而下划分为前端展示层")
    if arch_desc_idx < 0:
        arch_desc_idx = idx + 1

    # Also update the architecture description text
    p = doc.paragraphs[arch_desc_idx]
    if "自上而下划分为" in p.text:
        new_text = (
            "本文设计的智能运动训练系统采用前后端分离的分层架构，整体结构如下图所示。"
            "系统划分为前端展示层、后端服务层、数据存储层与模型服务层四个层级，"
            "其中后端服务层作为系统的中心枢纽，"
            "向上通过REST API向前端提供服务，向下分别调用数据存储层和模型服务层。"
            "数据存储层与模型服务层之间不存在直接交互，所有数据流均经由后端服务层统一调度。"
            "后端服务层内部按功能划分为RAG检索模块、长期记忆模块和多智能体模块三大核心模块，"
            "分别对应第1.5节中提出的三类关键技术方案。"
        )
        replace_paragraph_text(p, new_text)
        print(f"  已更新架构描述文字 (段落 {arch_desc_idx})")

    # Insert new architecture diagram image
    ref = doc.paragraphs[arch_desc_idx]
    caption = insert_paragraph_after(doc, ref, "图3-1 系统整体架构图", "Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.size = Pt(9)

    img_para = insert_paragraph_after(doc, ref, "", "Normal")
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Cm(14))

    print(f"  已插入新架构图")


# ═══════════════════════════════════════════════════════════
# Phase 3: Structural Changes
# ═══════════════════════════════════════════════════════════

def task_3_1_move_tech_stack(doc):
    """意见3: 将技术栈从第四章移至第三章"""
    print("\n[Task 3.1] 意见3: 技术栈移至设计章节")

    # Find Ch4 heading "开发环境与技术栈" (must be a heading style, not body text)
    ch4_tech_idx = -1
    toc_end = get_toc_end(doc)
    for i, p in enumerate(doc.paragraphs):
        if i <= toc_end:
            continue
        if "开发环境与技术栈" in p.text and "标题" in (p.style.name or ""):
            ch4_tech_idx = i
            break
    if ch4_tech_idx < 0:
        print("  [ERROR] 未找到4.1节标题")
        return

    # Rename Ch4 heading
    p = doc.paragraphs[ch4_tech_idx]
    replace_paragraph_text(p, p.text.replace("开发环境与技术栈", "开发环境配置"))
    print(f"  已将4.1标题改为'开发环境配置' (段落 {ch4_tech_idx})")

    # Find Ch3 "本章小结" to insert tech selection section before it
    ch3_summary_idx = find_body_paragraph_index(doc, "本章从需求分析出发")
    if ch3_summary_idx < 0:
        ch3_summary_idx = find_body_paragraph_index(doc, "本章小结", start=350)
    if ch3_summary_idx < 0:
        print("  [WARN] 未找到第三章本章小结，跳过技术栈移动")
        return

    # Insert new "技术选型" section before Ch3 summary
    ref = doc.paragraphs[ch3_summary_idx]

    tech_paragraphs = [
        ("技术选型", "2级标题"),

        ("系统的技术选型需要在开发效率、运行性能和生态兼容性之间取得平衡。"
         "本节从前端框架、后端框架、向量数据库、关系数据库和多智能体编排框架五个方面说明选型依据。", "Normal"),

        ("前端选用Vue3框架配合Vite构建工具。Vue3的组合式API适合组织训练计划管理、"
         "记忆查看等复杂交互页面的状态逻辑；Vite提供毫秒级的热更新体验，"
         "有利于快速迭代前端功能。UI组件库选用Element Plus，可减少基础组件的重复开发。", "Normal"),

        ("后端选用FastAPI框架。FastAPI原生支持异步请求处理，"
         "适合多智能体并行推理场景下的高并发需求；"
         "其自动生成的OpenAPI文档便于前后端协作调试。"
         "Python生态与LangChain、LangGraph等AI开发框架的兼容性也是选择FastAPI的重要原因。", "Normal"),

        ("向量数据库选用ChromaDB。相比Milvus或Weaviate等分布式方案，"
         "ChromaDB以嵌入式方式运行，无需独立部署集群，"
         "适合本系统当前万级向量规模的知识库体量。"
         "ChromaDB默认采用HNSW索引与余弦相似度度量，"
         "能够在当前数据规模下实现毫秒级的语义检索响应。"
         "此外，ChromaDB与LangChain生态深度集成，可直接作为检索器使用。", "Normal"),

        ("关系数据库选用SQLite。系统当前为单机部署的原型系统，"
         "用户规模在百人量级，SQLite的零配置、嵌入式特性可显著降低部署复杂度。"
         "SQLite支持完整的SQL语法和事务机制，能够满足用户数据、训练记录和记忆信息的持久化需求。"
         "若后续系统扩展至多实例部署，可替换为PostgreSQL，"
         "由于数据访问层已通过统一接口封装，迁移成本可控。", "Normal"),

        ("多智能体编排框架选用LangGraph。LangGraph以状态图描述应用逻辑，"
         "支持显式的状态管理、条件分支与循环边，"
         "相比LangChain传统的线性链式调用，更适合本系统中多教练协同推理的控制流需求。"
         "与AutoGen和MetaGPT等框架相比，"
         "LangGraph强调以有限状态加明确转移为基础的可控编排，"
         "与本系统对可维护性、可扩展性与异常降级的需求高度契合。", "Normal"),
    ]

    for text, style in reversed(tech_paragraphs):
        insert_paragraph_before(doc, ref, text, style)

    print(f"  已在第三章本章小结前插入技术选型节 ({len(tech_paragraphs)} 段)")


# ═══════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("论文修改脚本 v2")
    print(f"输入: {INPUT_FILE}")
    print(f"输出: {OUTPUT_FILE}")
    print("=" * 60)

    if not INPUT_FILE.exists():
        print(f"[ERROR] 输入文件不存在: {INPUT_FILE}")
        return

    # Copy input to output first
    shutil.copy2(INPUT_FILE, OUTPUT_FILE)
    doc = Document(str(OUTPUT_FILE))
    print(f"\n文档加载成功，共 {len(doc.paragraphs)} 个段落\n")

    # ── Phase 0: Generate diagrams ──
    print("=" * 40)
    print("Phase 0: 生成图表")
    print("=" * 40)
    use_case_img = generate_use_case_diagram()
    er_img = generate_er_diagram()
    arch_img = generate_architecture_diagram()

    # ── Phase 1: Text modifications ──
    print("\n" + "=" * 40)
    print("Phase 1: 纯文本修改")
    print("=" * 40)
    task_1_1_toc_cleanup(doc)
    task_1_2_distinguish_retain_utilize(doc)
    task_1_3_rewrite_problem2(doc)
    task_1_4_rewrite_section_1_5(doc)

    # ── Phase 2: Content insertions & diagrams ──
    print("\n" + "=" * 40)
    print("Phase 2: 内容插入与图表")
    print("=" * 40)
    task_2_1_add_rag_pipeline_description(doc)
    task_2_2_add_multiagent_detail(doc)
    task_2_3_insert_use_case_diagram(doc, use_case_img)
    task_2_4_insert_er_diagram(doc, er_img)
    task_2_5_replace_architecture_diagram(doc, arch_img)

    # ── Phase 3: Structural changes ──
    print("\n" + "=" * 40)
    print("Phase 3: 结构调整")
    print("=" * 40)
    task_3_1_move_tech_stack(doc)

    # ── Save ──
    doc.save(str(OUTPUT_FILE))
    print("\n" + "=" * 60)
    print(f"[DONE] 修改完成，已保存到: {OUTPUT_FILE}")
    print("=" * 60)


if __name__ == "__main__":
    main()
