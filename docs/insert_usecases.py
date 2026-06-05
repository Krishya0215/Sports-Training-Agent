"""
Insert use case diagrams and descriptions into the thesis document.
Each diagram is inserted after the corresponding functional requirement paragraph,
followed by a figure caption and a use case description table.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import copy

DOC_PATH = '22301126-刘冰彦-基于多智能体与检索增强生成的智能运动训练系统.docx'
OUTPUT_PATH = '22301126-刘冰彦-基于多智能体与检索增强生成的智能运动训练系统.docx'
DIAGRAM_DIR = 'use_case_diagrams'

doc = Document(DOC_PATH)

FIGURE_START_NUM = 1

use_case_data = [
    {
        'marker': '（1）用户信息管理功能',
        'image': 'usecase_1_user_info.png',
        'caption': '用户信息管理用例图',
        'description': (
            '如图3-x所示，用户信息管理用例图描述了用户与系统在账户及个人资料管理方面的交互。'
            '参与者为"用户"，包含5个用例：用户注册、用户登录、找回密码、查看个人资料和编辑个人资料。'
            '其中，编辑个人资料以<<include>>关系包含查看个人资料，即用户在编辑前需先查看当前资料内容。'
            '用户注册用例要求用户填写基本身体信息（身高、体重、年龄、性别等）与训练相关信息（训练目标、训练经验、伤病史等），'
            '这些数据将作为后续个性化训练指导的基础输入，同时也是语义记忆层建立用户长期画像的数据源头。'
        ),
    },
    {
        'marker': '（2）AI教练问答功能',
        'image': 'usecase_2_ai_coach.png',
        'caption': 'AI教练问答用例图',
        'description': (
            '如图3-x所示，AI教练问答用例图描述了用户与系统在智能问答方面的核心交互。'
            '参与者为"用户"，包含7个用例：提出训练问题、选择问答模式、单智能体问答、多智能体协同问答、'
            '查看引用来源、查看对话历史和意图识别与路由。'
            '其中，提出训练问题以<<include>>关系包含选择问答模式；选择问答模式以<<include>>关系包含单智能体问答，'
            '以<<extend>>关系扩展至多智能体协同问答，即当用户问题涉及多个专业领域时，系统扩展激活多智能体协同模式；'
            '多智能体协同问答以<<include>>关系包含意图识别与路由，由系统自动完成多标签意图分类与教练调度。'
            '该模块是系统最主要的交互入口，RAG检索与引用机制嵌入问答链路，确保每条回答具备可追溯的文献依据。'
        ),
    },
    {
        'marker': '（3）训练计划功能',
        'image': 'usecase_3_training_plan.png',
        'caption': '训练计划管理用例图',
        'description': (
            '如图3-x所示，训练计划管理用例图描述了用户与系统在训练计划生命周期管理方面的交互。'
            '参与者为"用户"，包含5个用例：生成训练计划、查看训练计划、修改训练计划、归档训练计划和读取用户画像。'
            '其中，生成训练计划以<<include>>关系包含读取用户画像，即系统在生成计划前需自动从长期记忆中提取用户的'
            '训练目标、当前水平、可用时间、伤病史等个性化约束条件。'
            '生成的训练计划包括训练目标、周期安排与具体训练内容，将长期记忆中沉淀的用户状态显式落成结构化、可追踪的训练方案，'
            '用户可在界面上查看、修改与归档，是系统提供长期性训练管理的直观体现。'
        ),
    },
    {
        'marker': '（4）健康记录功能',
        'image': 'usecase_4_health_record.png',
        'caption': '健康记录管理用例图',
        'description': (
            '如图3-x所示，健康记录管理用例图描述了用户与系统在持续健康数据采集方面的交互。'
            '参与者为"用户"，包含5个用例：记录训练表现、记录每日饮食、记录体重数据、查看健康趋势和写入情景记忆。'
            '其中，记录训练表现、记录每日饮食和记录体重数据三个用例均以<<include>>关系包含写入情景记忆，'
            '即用户每次提交健康数据后，系统自动将其作为事件级信息写入长期记忆的情景层。'
            '这些持续累积的时序数据为情景记忆提供了真实的事件来源，使系统能够动态感知用户近期身体状况变化，'
            '从而在问答与计划生成中给出贴合实际的个性化判断。'
        ),
    },
    {
        'marker': '（5）知识库管理功能',
        'image': 'usecase_5_knowledge_base.png',
        'caption': '知识库管理用例图',
        'description': (
            '如图3-x所示，知识库管理用例图描述了管理员与系统在知识库维护方面的交互。'
            '参与者为"管理员"，包含5个用例：上传训练文档、预览文档内容、删除知识文档、文档解析与切分、向量化与入库。'
            '其中，上传训练文档以<<include>>关系包含文档解析与切分，文档解析与切分又以<<include>>关系包含向量化与入库，'
            '形成"上传-解析-向量化"的自动化处理链路。'
            '管理员上传运动训练领域的专业文档后，系统自动完成文档解析、语义切分、嵌入向量化与ChromaDB入库，'
            '为RAG检索提供持续更新的权威知识基础，确保训练建议的科学性与时效性。'
        ),
    },
    {
        'marker': '（6）记忆管理功能',
        'image': 'usecase_6_memory_mgmt.png',
        'caption': '记忆管理用例图',
        'description': (
            '如图3-x所示，记忆管理用例图描述了系统内部在记忆生命周期管理方面的自动化交互。'
            '参与者为"系统"，包含7个用例：维护工作记忆、记录情景记忆、提取语义记忆、重要性评分、'
            '记忆遗忘与衰减、记忆感知检索和跨会话状态追踪。'
            '其中，记录情景记忆和提取语义记忆均以<<include>>关系包含重要性评分，即系统在写入长期记忆前'
            '需先评估信息价值，只有超过阈值的内容才进入持久化存储；记忆感知检索以<<include>>关系包含跨会话状态追踪，'
            '即检索过程中自动融合用户长期状态信息。'
            '该模块以三层记忆机制为技术基座，对内承担用户状态的持续追踪与维护，'
            '对外为AI教练问答、训练计划等功能提供统一的用户状态信号，是实现系统"记得住用户"的关键支撑。'
        ),
    },
]


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    run = new_para.add_run(text)
    return new_para


def insert_image_paragraph_after(paragraph, image_path, width_inches=4.8):
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return new_para


def set_paragraph_font(paragraph, font_name='宋体', font_size=Pt(12), first_line_indent=None):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
    if first_line_indent is not None:
        pf = paragraph.paragraph_format
        pf.first_line_indent = first_line_indent


fig_counter = 0

for uc in reversed(use_case_data):
    marker = uc['marker']
    target_para = None
    target_idx = None

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(marker):
            target_para = p
            target_idx = i
            break

    if target_para is None:
        print(f"WARNING: Could not find paragraph starting with '{marker}'")
        continue

    print(f"Found '{marker}' at paragraph {target_idx}")

    desc_para = insert_paragraph_after(target_para, uc['description'])
    desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_font(desc_para, font_name='宋体', font_size=Pt(12), first_line_indent=Cm(0.74))

    caption_para = insert_paragraph_after(target_para, f"图3-x {uc['caption']}")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(caption_para, font_name='黑体', font_size=Pt(10.5))

    image_path = os.path.join(DIAGRAM_DIR, uc['image'])
    img_para = insert_image_paragraph_after(target_para, image_path, width_inches=4.5)

doc.save(OUTPUT_PATH)
print(f"\nDocument saved to {OUTPUT_PATH}")
print("NOTE: Please manually update figure numbers (3-x) in the document.")
