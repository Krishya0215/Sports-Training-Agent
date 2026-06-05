"""
Renumber figures and tables in chapter 3 after inserting 6 use case diagrams.
- New diagrams: 图3-1 through 图3-6
- Existing 图3-1 becomes 图3-7, 图3-2 becomes 图3-8, etc.
- Existing 表3-1 becomes 表3-7, 表3-2 becomes 表3-8, etc.
"""
from docx import Document
import re

DOC_PATH = '22301126-刘冰彦-基于多智能体与检索增强生成的智能运动训练系统.docx'

doc = Document(DOC_PATH)

fig_num = 1
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if '图3-x' in text:
        old_text = '图3-x'
        new_text = f'图3-{fig_num}'
        for run in p.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
        print(f'[{i}] Renamed caption: {new_text}')
        fig_num += 1

OFFSET = 6

for i, p in enumerate(doc.paragraphs):
    text = p.text
    changed = False
    for run in p.runs:
        original = run.text
        new = original
        for old_num in range(10, 0, -1):
            new_num = old_num + OFFSET
            new = new.replace(f'图3-{old_num}', f'图3-{new_num}')
            new = new.replace(f'表3-{old_num}', f'表3-{new_num}')
        if new != original:
            run.text = new
            changed = True
    if changed:
        print(f'[{i}] Updated: {p.text.strip()[:80]}')

doc.save(DOC_PATH)
print(f"\nDocument saved. Figures renumbered with offset +{OFFSET}.")
